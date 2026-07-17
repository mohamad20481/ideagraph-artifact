"""
growth.py - User growth, retention, and viral features for IdeaGraph.

Implements 5 features for reaching 1M users:

  1. WeeklyChallengeSystem  - time-limited goals with XP rewards
  2. TrendingFeed           - public page showing top ideas across all users
  3. ResearchProposalExport - one-click 2-page grant-ready proposal PDF
  4. IdeaEvolutionTree      - visual family tree of idea lineage
  5. ReferralProgram        - invite code → bonus for both parties
"""

from __future__ import annotations

import hashlib
import json
import os
import random
import secrets
import time
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

try:
    import db
except ImportError:
    db = None  # type: ignore


# ─────────────────────────────────────────────────────────────────────────────
# 1. WEEKLY CHALLENGE SYSTEM
# ─────────────────────────────────────────────────────────────────────────────

# Challenge templates — rotated weekly. Each has a goal, XP reward, and check fn.
CHALLENGE_TEMPLATES = [
    {
        "id": "generate_3",
        "title": "Idea Sprint",
        "description": "Generate at least 3 research ideas this week",
        "goal_type": "ideas_generated",
        "goal_count": 3,
        "xp_reward": 75,
        "badge": "sprint_winner",
    },
    {
        "id": "high_quality_1",
        "title": "Quality Quest",
        "description": "Generate 1 idea with quality score above 0.7",
        "goal_type": "high_quality_idea",
        "goal_count": 1,
        "quality_threshold": 0.7,
        "xp_reward": 100,
        "badge": "quality_seeker",
    },
    {
        "id": "share_2",
        "title": "Community Contributor",
        "description": "Share 2 ideas publicly this week",
        "goal_type": "ideas_shared",
        "goal_count": 2,
        "xp_reward": 60,
        "badge": "contributor",
    },
    {
        "id": "explore_new_method",
        "title": "Methodology Explorer",
        "description": "Generate ideas using 3 different methodology types",
        "goal_type": "unique_methodologies",
        "goal_count": 3,
        "xp_reward": 80,
        "badge": "explorer",
    },
    {
        "id": "bookmark_5",
        "title": "Curator",
        "description": "Bookmark 5 ideas you find promising",
        "goal_type": "bookmarks_added",
        "goal_count": 5,
        "xp_reward": 50,
        "badge": "curator",
    },
    {
        "id": "run_2_topics",
        "title": "Cross-Domain Thinker",
        "description": "Run the pipeline on 2 different topics this week",
        "goal_type": "unique_topics",
        "goal_count": 2,
        "xp_reward": 90,
        "badge": "cross_domain",
    },
    {
        "id": "streak_5",
        "title": "Consistency Champion",
        "description": "Log in 5 days in a row this week",
        "goal_type": "login_streak",
        "goal_count": 5,
        "xp_reward": 70,
        "badge": "consistent",
    },
]


def get_current_challenge() -> Dict[str, Any]:
    """
    Return the current week's challenge. Rotates deterministically based on
    ISO week number so all users see the same challenge.
    """
    now = datetime.now()
    week_num = now.isocalendar()[1]
    idx = week_num % len(CHALLENGE_TEMPLATES)
    challenge = dict(CHALLENGE_TEMPLATES[idx])
    # Add timing
    monday = now - timedelta(days=now.weekday())
    monday = monday.replace(hour=0, minute=0, second=0, microsecond=0)
    sunday = monday + timedelta(days=6, hours=23, minutes=59, seconds=59)
    challenge["week_start"] = monday.isoformat()
    challenge["week_end"] = sunday.isoformat()
    challenge["days_remaining"] = max(0, (sunday - now).days)
    return challenge


def get_challenge_progress(user_id: int) -> Dict[str, Any]:
    """
    Check the user's progress on the current week's challenge.
    Returns {challenge: {...}, progress: int, completed: bool, claimed: bool}
    """
    challenge = get_current_challenge()
    week_start = challenge["week_start"][:10]

    progress = 0
    if not db:
        return {"challenge": challenge, "progress": 0, "completed": False, "claimed": False}

    try:
        with db._lock:
            conn = db._get_conn()
            try:
                goal = challenge["goal_type"]
                if goal == "ideas_generated":
                    row = conn.execute(
                        "SELECT COALESCE(SUM(ideas_count), 0) FROM results "
                        "WHERE user_id = ? AND created_at >= ?",
                        (user_id, week_start),
                    ).fetchone()
                    progress = row[0] if row else 0
                elif goal == "high_quality_idea":
                    # Push JSON aggregation into SQL via SQLite's JSON1 extension
                    # so we don't ship hundreds of KB of `results_json` to Python
                    # and re-parse every row. Falls back to the Python loop only
                    # if a row's JSON is malformed (json_each raises) or JSON1
                    # isn't compiled in.
                    threshold = challenge.get("quality_threshold", 0.7)
                    try:
                        row = conn.execute(
                            "SELECT COUNT(*) FROM results r, "
                            "json_each(r.results_json, '$.ideas') AS i "
                            "WHERE r.user_id = ? AND r.created_at >= ? "
                            "AND CAST(json_extract(i.value, '$.quality_score') AS REAL) >= ?",
                            (user_id, week_start, threshold),
                        ).fetchone()
                        progress = row[0] if row else 0
                    except Exception:
                        rows = conn.execute(
                            "SELECT results_json FROM results "
                            "WHERE user_id = ? AND created_at >= ?",
                            (user_id, week_start),
                        ).fetchall()
                        for r in rows:
                            try:
                                data = json.loads(r["results_json"])
                            except Exception:
                                continue
                            for idea in data.get("ideas", []):
                                if idea.get("quality_score", 0) >= threshold:
                                    progress += 1
                elif goal == "ideas_shared":
                    row = conn.execute(
                        "SELECT COUNT(*) FROM share_tokens "
                        "WHERE user_id = ? AND created_at >= ?",
                        (user_id, week_start),
                    ).fetchone()
                    progress = row[0] if row else 0
                elif goal == "unique_methodologies":
                    # Same SQL-side aggregation: SQLite walks the JSON natively
                    # in C and DISTINCTs in-engine, so Python only sees the
                    # final unique-count. Falls back to a Python set if JSON1
                    # is unavailable or any row is malformed.
                    try:
                        row = conn.execute(
                            "SELECT COUNT(DISTINCT json_extract(i.value, '$.methodology_type')) "
                            "FROM results r, json_each(r.results_json, '$.ideas') AS i "
                            "WHERE r.user_id = ? AND r.created_at >= ? "
                            "AND json_extract(i.value, '$.methodology_type') IS NOT NULL",
                            (user_id, week_start),
                        ).fetchone()
                        progress = row[0] if row else 0
                    except Exception:
                        rows = conn.execute(
                            "SELECT results_json FROM results "
                            "WHERE user_id = ? AND created_at >= ?",
                            (user_id, week_start),
                        ).fetchall()
                        methods = set()
                        for r in rows:
                            try:
                                data = json.loads(r["results_json"])
                            except Exception:
                                continue
                            for idea in data.get("ideas", []):
                                m = idea.get("methodology_type")
                                if m:
                                    methods.add(m)
                        progress = len(methods)
                elif goal == "bookmarks_added":
                    row = conn.execute(
                        "SELECT COUNT(*) FROM bookmarks "
                        "WHERE user_id = ? AND created_at >= ?",
                        (user_id, week_start),
                    ).fetchone()
                    progress = row[0] if row else 0
                elif goal == "unique_topics":
                    row = conn.execute(
                        "SELECT COUNT(DISTINCT topic) FROM results "
                        "WHERE user_id = ? AND created_at >= ?",
                        (user_id, week_start),
                    ).fetchone()
                    progress = row[0] if row else 0
                elif goal == "login_streak":
                    row = conn.execute(
                        "SELECT current_streak FROM user_stats WHERE user_id = ?",
                        (user_id,),
                    ).fetchone()
                    progress = row["current_streak"] if row else 0

                # Check if already claimed this week
                claimed = False
                claim_row = conn.execute(
                    "SELECT id FROM achievements WHERE user_id = ? AND badge_key = ?",
                    (user_id, f"challenge_{challenge['id']}_{week_start}"),
                ).fetchone()
                if claim_row:
                    claimed = True

            finally:
                conn.close()
    except Exception:
        pass

    completed = progress >= challenge["goal_count"]
    return {
        "challenge": challenge,
        "progress": min(progress, challenge["goal_count"]),
        "completed": completed,
        "claimed": claimed,
    }


def claim_challenge_reward(user_id: int) -> Dict[str, Any]:
    """Claim XP reward for completing the weekly challenge."""
    status = get_challenge_progress(user_id)
    if not status["completed"]:
        return {"success": False, "reason": "Challenge not completed yet"}
    if status["claimed"]:
        return {"success": False, "reason": "Already claimed this week"}

    challenge = status["challenge"]
    week_start = challenge["week_start"][:10]
    xp = challenge["xp_reward"]

    try:
        import engagement
        engagement.award_xp(user_id, xp, reason=f"challenge:{challenge['id']}")
        # Record claim as achievement so it can't be claimed twice
        with db._lock:
            conn = db._get_conn()
            try:
                conn.execute(
                    "INSERT OR IGNORE INTO achievements (user_id, badge_key) VALUES (?, ?)",
                    (user_id, f"challenge_{challenge['id']}_{week_start}"),
                )
                conn.commit()
            finally:
                conn.close()
        return {"success": True, "xp_earned": xp, "badge": challenge.get("badge", "")}
    except Exception as e:
        return {"success": False, "reason": str(e)}


# ─────────────────────────────────────────────────────────────────────────────
# 2. PUBLIC TRENDING FEED
# ─────────────────────────────────────────────────────────────────────────────

def get_trending_ideas(limit: int = 20) -> List[Dict[str, Any]]:
    """
    Get the top publicly shared ideas ranked by a trending score:
      score = views + 3*likes + 10*recency_bonus

    recency_bonus = max(0, 7 - days_since_shared) / 7
    """
    if not db:
        return []
    try:
        with db._lock:
            conn = db._get_conn()
            try:
                rows = conn.execute("""
                    SELECT s.token, s.idea_json, s.topic, s.views, s.likes,
                           s.created_at, u.username
                    FROM share_tokens s
                    JOIN users u ON s.user_id = u.id
                    ORDER BY s.created_at DESC
                    LIMIT 200
                """).fetchall()
            finally:
                conn.close()

        now = datetime.now()
        scored = []
        for row in rows:
            try:
                idea = json.loads(row["idea_json"])
                created = datetime.fromisoformat(row["created_at"].replace("Z", ""))
                days_ago = (now - created).days
                recency = max(0, 7 - days_ago) / 7.0
                score = row["views"] + 3 * row["likes"] + 10 * recency
                scored.append({
                    "token": row["token"],
                    "title": idea.get("title", "Untitled"),
                    "method_preview": (idea.get("method", "") or "")[:120],
                    "quality_score": idea.get("quality_score", 0),
                    "methodology_type": idea.get("methodology_type", ""),
                    "topic": row["topic"],
                    "views": row["views"],
                    "likes": row["likes"],
                    "username": row["username"],
                    "created_at": row["created_at"],
                    "trending_score": round(score, 1),
                })
            except Exception:
                continue

        scored.sort(key=lambda x: x["trending_score"], reverse=True)
        return scored[:limit]
    except Exception:
        return []


def render_trending_feed(st_module) -> None:
    """Render the trending feed as a Streamlit component."""
    st = st_module
    trending = get_trending_ideas(limit=15)
    if not trending:
        st.info("No shared ideas yet. Be the first to share!")
        return

    st.markdown("### Trending Research Ideas")
    st.caption("The hottest ideas shared by the community this week")

    for rank, item in enumerate(trending, 1):
        medal = {1: "🥇", 2: "🥈", 3: "🥉"}.get(rank, f"**{rank}.**")
        q = item["quality_score"]
        q_color = "🟢" if q >= 0.6 else "🟡" if q >= 0.4 else "🔴"

        col1, col2 = st.columns([4, 1])
        with col1:
            st.markdown(
                f"{medal} **{item['title'][:60]}**\n\n"
                f"{q_color} q={q:.2f} | {item['methodology_type'].replace('_',' ').title()} | "
                f"by @{item['username']} | {item['views']} views, {item['likes']} likes"
            )
            if item["method_preview"]:
                st.caption(item["method_preview"])
        with col2:
            base = os.getenv("APP_BASE_URL", "http://localhost:8510")
            st.markdown(f"[View]({base}/?share={item['token']})")
        st.divider()


# ─────────────────────────────────────────────────────────────────────────────
# 3. RESEARCH PROPOSAL EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_proposal_markdown(
    idea: Dict[str, Any],
    dag_papers: List[Dict[str, Any]] = None,
    topic: str = "",
) -> str:
    """
    Generate a 2-page research proposal in Markdown from an idea.
    Includes: title, abstract, motivation, related work, method, expected
    results, timeline, budget estimate.
    """
    title = idea.get("title", "Untitled Research Proposal")
    motivation = idea.get("motivation", "")
    method = idea.get("method", "")
    hypothesis = idea.get("hypothesis", "")
    resources = idea.get("resources", "")
    expected_outcome = idea.get("expected_outcome", "")
    risk = idea.get("risk_assessment", "")
    quality = idea.get("quality_score", 0)

    # Build related work from DAG papers
    related_work = ""
    if dag_papers:
        related_work = "## Related Work\n\n"
        for p in dag_papers[:6]:
            p_title = p.get("title", "")
            p_year = p.get("year", "")
            p_abstract = (p.get("abstract", "") or "")[:150]
            if p_title:
                related_work += f"- **{p_title}** ({p_year}): {p_abstract}...\n"
        related_work += "\n"

    # Build timeline
    timeline = """## Timeline

| Phase | Duration | Deliverable |
|-------|----------|-------------|
| Literature review & data collection | Weeks 1-2 | Annotated bibliography, dataset ready |
| Method implementation | Weeks 3-5 | Working prototype, baseline results |
| Experiments & analysis | Weeks 6-8 | Full experimental results, ablations |
| Paper writing & revision | Weeks 9-10 | Submission-ready manuscript |
| Buffer & revisions | Weeks 11-12 | Final polished version |
"""

    # Build budget
    budget = """## Budget Estimate

| Item | Cost | Justification |
|------|------|---------------|
| GPU compute (cloud) | $500-2,000 | Training on A100 instances |
| Dataset access | $0-200 | Public benchmarks + optional licensed data |
| Software licenses | $0 | Open-source stack (PyTorch, etc.) |
| Publication fees | $0-500 | Open access APC if applicable |
| **Total** | **$500-2,700** | |
"""

    proposal = f"""# Research Proposal: {title}

**Topic:** {topic or 'Research Investigation'}
**Quality Score:** {quality:.2f}/1.00
**Date:** {datetime.now().strftime('%B %d, %Y')}

---

## Abstract

{motivation}

**Hypothesis:** {hypothesis}

## Motivation & Significance

{motivation}

This research addresses a critical gap in the current literature. {expected_outcome}

{related_work}## Proposed Method

{method}

### Key Innovation

The core novelty of this approach lies in the combination of techniques described above,
which has not been previously explored in the literature.

## Expected Results

{expected_outcome}

## Risk Assessment & Mitigation

{risk}

{timeline}
{budget}
## References

"""
    if dag_papers:
        for i, p in enumerate(dag_papers[:8], 1):
            p_title = p.get("title", "")
            p_year = p.get("year", "")
            authors = p.get("authors", "")
            if isinstance(authors, list):
                authors = ", ".join(a.get("name", "") for a in authors[:3])
            proposal += f"[{i}] {authors}. \"{p_title}\". {p_year}.\n\n"

    proposal += "\n---\n*Generated by IdeaGraph — AI-Powered Research Ideation*\n"
    return proposal


def export_proposal_docx(idea: Dict[str, Any], dag_papers: List = None, topic: str = "") -> Optional[bytes]:
    """Export a research proposal as DOCX bytes."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    md = generate_proposal_markdown(idea, dag_papers, topic)
    doc = Document()

    # Title
    title_para = doc.add_heading(idea.get("title", "Research Proposal"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown sections into DOCX
    current_section = ""
    for line in md.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            continue  # skip top title (already added)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("| "):
            # Simple table row (skip complex parsing, add as text)
            doc.add_paragraph(line, style="List Bullet")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        elif line.startswith("---"):
            doc.add_page_break()
        elif line:
            doc.add_paragraph(line)

    buf = __import__("io").BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# 4. IDEA EVOLUTION TREE
# ─────────────────────────────────────────────────────────────────────────────

def build_evolution_tree(ideas: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build a tree structure from ideas based on parent_title linkage.
    Returns a nested dict suitable for visualization.
    """
    # Index ideas by title
    by_title: Dict[str, Dict] = {}
    for idea in ideas:
        title = idea.get("title", "")
        if title:
            by_title[title] = {
                "title": title,
                "quality_score": idea.get("quality_score", 0),
                "generation": idea.get("generation", 0),
                "methodology_type": idea.get("methodology_type", ""),
                "novelty_level": idea.get("novelty_level", ""),
                "children": [],
            }

    # Link parents → children
    roots = []
    for idea in ideas:
        title = idea.get("title", "")
        parent = idea.get("parent_title", "")
        node = by_title.get(title)
        if not node:
            continue
        if parent and parent in by_title:
            by_title[parent]["children"].append(node)
        else:
            roots.append(node)

    return {"roots": roots, "total_ideas": len(ideas), "total_roots": len(roots)}


def render_evolution_tree_plotly(ideas: List[Dict[str, Any]]):
    """
    Render idea evolution as an interactive Plotly treemap.
    Color = quality, size = method length (specificity proxy).
    """
    try:
        import plotly.express as px
    except ImportError:
        return None

    if not ideas:
        return None

    # Flatten for treemap: each idea needs parent reference
    labels = []
    parents = []
    values = []
    colors = []
    hovers = []

    # Add root node
    labels.append("All Ideas")
    parents.append("")
    values.append(0)
    colors.append(0.5)
    hovers.append("Root of idea evolution tree")

    for idea in ideas:
        title = idea.get("title", "Untitled")[:40]
        parent = idea.get("parent_title", "")
        q = idea.get("quality_score", 0)
        gen = idea.get("generation", 0)
        method_len = len(idea.get("method", ""))

        labels.append(title)
        # If parent exists in our list, link to it; otherwise link to root
        if parent and any(i.get("title", "")[:40] == parent[:40] for i in ideas):
            parents.append(parent[:40])
        else:
            parents.append("All Ideas")
        values.append(max(method_len, 50))
        colors.append(q)
        hovers.append(
            f"Quality: {q:.2f}<br>"
            f"Generation: {gen}<br>"
            f"Method: {idea.get('methodology_type', '?')}<br>"
            f"Novelty: {idea.get('novelty_level', '?')}"
        )

    fig = px.treemap(
        names=labels,
        parents=parents,
        values=values,
        color=colors,
        color_continuous_scale="RdYlGn",
        range_color=[0, 1],
        hover_data={"hover": hovers} if hovers else None,
    )
    fig.update_layout(
        title="Idea Evolution Tree (color = quality, size = specificity)",
        template="plotly_dark",
        height=500,
        margin=dict(t=50, l=10, r=10, b=10),
    )
    fig.update_traces(
        hovertemplate="<b>%{label}</b><br>%{customdata[0]}<extra></extra>",
        customdata=[[h] for h in hovers],
    )
    return fig


# ─────────────────────────────────────────────────────────────────────────────
# 5. REFERRAL PROGRAM
# ─────────────────────────────────────────────────────────────────────────────

def generate_referral_code(user_id: int) -> str:
    """Generate a deterministic but unguessable referral code for a user."""
    raw = f"ideagraph-ref-{user_id}-{os.getenv('REFERRAL_SALT', 'default-salt')}"
    return hashlib.sha256(raw.encode()).hexdigest()[:8].upper()


def get_or_create_referral(user_id: int) -> Dict[str, Any]:
    """Get referral stats or create a new referral code."""
    code = generate_referral_code(user_id)
    if not db:
        return {"code": code, "referrals": 0, "xp_earned": 0}

    try:
        with db._lock:
            conn = db._get_conn()
            try:
                # Ensure referrals table exists (migration-safe)
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS referrals (
                        id          INTEGER PRIMARY KEY AUTOINCREMENT,
                        referrer_id INTEGER NOT NULL,
                        referred_id INTEGER NOT NULL,
                        code        TEXT    NOT NULL,
                        xp_awarded  INTEGER NOT NULL DEFAULT 0,
                        created_at  TEXT    NOT NULL DEFAULT (datetime('now')),
                        FOREIGN KEY (referrer_id) REFERENCES users(id),
                        FOREIGN KEY (referred_id) REFERENCES users(id),
                        UNIQUE(referred_id)
                    )
                """)
                conn.commit()

                row = conn.execute(
                    "SELECT COUNT(*) as cnt, COALESCE(SUM(xp_awarded), 0) as xp "
                    "FROM referrals WHERE referrer_id = ?",
                    (user_id,),
                ).fetchone()
                return {
                    "code": code,
                    "referrals": row["cnt"] if row else 0,
                    "xp_earned": row["xp"] if row else 0,
                }
            finally:
                conn.close()
    except Exception:
        return {"code": code, "referrals": 0, "xp_earned": 0}


def apply_referral(referred_user_id: int, referral_code: str) -> Dict[str, Any]:
    """
    Apply a referral code during/after registration.
    Awards XP to both referrer and referred user.
    """
    if not db:
        return {"success": False, "reason": "DB not available"}

    REFERRER_XP = 100
    REFERRED_XP = 50

    try:
        # Find the referrer by checking all users' codes
        with db._lock:
            conn = db._get_conn()
            try:
                users = conn.execute("SELECT id FROM users").fetchall()
                referrer_id = None
                for u in users:
                    if generate_referral_code(u["id"]) == referral_code.upper():
                        referrer_id = u["id"]
                        break

                if not referrer_id:
                    return {"success": False, "reason": "Invalid referral code"}
                if referrer_id == referred_user_id:
                    return {"success": False, "reason": "Cannot refer yourself"}

                # Check if already referred
                existing = conn.execute(
                    "SELECT id FROM referrals WHERE referred_id = ?",
                    (referred_user_id,),
                ).fetchone()
                if existing:
                    return {"success": False, "reason": "Already used a referral code"}

                # Record referral
                conn.execute(
                    "INSERT INTO referrals (referrer_id, referred_id, code, xp_awarded) "
                    "VALUES (?, ?, ?, ?)",
                    (referrer_id, referred_user_id, referral_code, REFERRER_XP),
                )
                conn.commit()
            finally:
                conn.close()

        # Award XP to both
        try:
            import engagement
            engagement.award_xp(referrer_id, REFERRER_XP, reason="referral_given")
            engagement.award_xp(referred_user_id, REFERRED_XP, reason="referral_received")
        except Exception:
            pass

        return {
            "success": True,
            "referrer_xp": REFERRER_XP,
            "referred_xp": REFERRED_XP,
        }
    except Exception as e:
        return {"success": False, "reason": str(e)}


def get_referral_link(user_id: int) -> str:
    """Get the full referral URL for a user."""
    code = generate_referral_code(user_id)
    base = os.getenv("APP_BASE_URL", "http://localhost:8510")
    return f"{base}/?ref={code}"
